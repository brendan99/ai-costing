import streamlit as st
import chromadb
import os
from pathlib import Path
import json
import requests
from typing import List, Dict, Any
import time
import hashlib
from datetime import datetime
import mimetypes
from pypdf import PdfReader
from docx import Document as DocxDocument
import gc
import shutil
import psutil
from sentence_transformers import SentenceTransformer
from legal_chunking import legal_aware_chunk_text
from legal_entity_extraction import extract_entities
from collections import defaultdict
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Constants
UPLOAD_DIR = "./uploads"
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 128))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 50))
MAX_RETRIES = 3
RETRY_DELAY = 1  # seconds
MAX_DOCUMENT_SIZE = int(os.getenv("MAX_DOCUMENT_SIZE", "10485760"))  # 10MB in bytes
PROCESSED_DIR = "./processed"
INDEX_LOG = os.path.join(PROCESSED_DIR, "indexed_files.json")

# Initialize session state
if 'documents_loaded' not in st.session_state:
    st.session_state.documents_loaded = False
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = []
if 'last_processed_hash' not in st.session_state:
    st.session_state.last_processed_hash = None
if 'processing_error' not in st.session_state:
    st.session_state.processing_error = None

# Add a global log buffer for recent activity
if 'log_buffer' not in st.session_state:
    st.session_state.log_buffer = []
if 'doc_status' not in st.session_state:
    st.session_state.doc_status = {}
if 'show_logs' not in st.session_state:
    st.session_state.show_logs = False
if 'processing' not in st.session_state:
    st.session_state.processing = False
if 'llm_responses' not in st.session_state:
    st.session_state.llm_responses = []

LOG_BUFFER_SIZE = 50

print("[LOG] Current working directory:", os.getcwd())

def init_storage():
    """Initialize storage directory and ensure it exists."""
    try:
        # Create all necessary directories with proper permissions
        for dir_path in ['storage', 'uploads', PROCESSED_DIR]:
            path = Path(dir_path)
            if path.exists():
                # Ensure directory is writable
                os.chmod(path, 0o777)
                # Clean up any SQLite files
                for file in path.glob("*.sqlite3"):
                    try:
                        os.chmod(file, 0o666)
                        os.remove(file)
                    except Exception as e:
                        print(f"Error removing SQLite file {file}: {e}")
            else:
                # Create new directory with proper permissions
                path.mkdir(parents=True, exist_ok=True)
                os.chmod(path, 0o777)
        
        return Path('storage')
    except Exception as e:
        st.error(f"Error initializing storage: {str(e)}")
        return None

# Call this at startup
init_storage()

def extract_document_metadata(file_path: str) -> dict:
    """Extract metadata from a document file."""
    metadata = {
        'filename': os.path.basename(file_path),
        'file_type': os.path.splitext(file_path)[1].lower(),
        'file_size': os.path.getsize(file_path),
        'created_date': datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
        'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
        'mime_type': mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
    }
    
    if file_path.lower().endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                pdf = PdfReader(f)
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('/Title', ''),
                        'author': pdf.metadata.get('/Author', ''),
                        'subject': pdf.metadata.get('/Subject', ''),
                        'keywords': pdf.metadata.get('/Keywords', ''),
                        'creator': pdf.metadata.get('/Creator', ''),
                        'producer': pdf.metadata.get('/Producer', ''),
                        'page_count': len(pdf.pages)
                    })
        except Exception as e:
            st.warning(f"Could not extract PDF metadata: {str(e)}")
            
    elif file_path.lower().endswith('.docx'):
        try:
            doc = DocxDocument(file_path)
            metadata.update({
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'keywords': doc.core_properties.keywords or '',
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else '',
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else '',
                'paragraph_count': len(doc.paragraphs)
            })
        except Exception as e:
            st.warning(f"Could not extract DOCX metadata: {str(e)}")
    
    return metadata

def get_directory_hash():
    """Calculate a hash of the uploads directory contents."""
    if not os.path.exists('uploads'):
        return None
    
    hash_md5 = hashlib.md5()
    for root, dirs, files in os.walk('uploads'):
        for file in sorted(files):
            file_path = os.path.join(root, file)
            with open(file_path, 'rb') as f:
                hash_md5.update(f.read())
    return hash_md5.hexdigest()

def check_documents_changed():
    """Check if documents have changed since last processing."""
    current_hash = get_directory_hash()
    if current_hash != st.session_state.last_processed_hash:
        st.session_state.last_processed_hash = current_hash
        return True
    return False

def log_memory_usage(note=""):
    process = psutil.Process(os.getpid())
    mem = process.memory_info().rss / (1024 * 1024 * 1024)  # in GB
    msg = f"[MEMORY] {note} - Current usage: {mem:.2f} GB"
    print(msg)
    add_log(msg)

def process_document(file_path: str) -> List[Dict[str, Any]]:
    """Process a document and return all chunks with metadata."""
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_DOCUMENT_SIZE:
            raise ValueError(f"File size {file_size} exceeds maximum allowed size of {MAX_DOCUMENT_SIZE} bytes")

        metadata = extract_document_metadata(file_path)
        chunks = []
        ext = os.path.splitext(file_path)[1].lower()
        text = ""

        add_log(f"Processing {os.path.basename(file_path)}")

        if ext == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    pdf = PdfReader(f)
                    for page in pdf.pages:
                        text += page.extract_text() or ""
            except Exception as e:
                st.error(f"Error reading PDF: {str(e)}")
                return []
        elif ext == ".docx":
            try:
                doc = DocxDocument(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except Exception as e:
                st.error(f"Error reading DOCX: {str(e)}")
                return []
        elif ext in [".txt", ".md"]:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except Exception as e:
                st.error(f"Error reading text file: {str(e)}")
                return []
        else:
            st.error(f"Unsupported file type: {ext}")
            return []

        # Use legal-aware chunking
        text_chunks = legal_aware_chunk_text(text, CHUNK_SIZE, CHUNK_OVERLAP)
        add_log(f"Created {len(text_chunks)} legal-aware chunks for {os.path.basename(file_path)}")
        
        for i, chunk in enumerate(text_chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'chunk_id': i,
                'total_chunks': len(text_chunks)
            })
            chunks.append({
                'text': chunk,
                'metadata': chunk_metadata
            })
        
        return chunks
    except Exception as e:
        st.error(f"Error processing document {file_path}: {str(e)}")
        print(f"Error processing document {file_path}: {str(e)}")
        return []

def save_uploaded_file(uploaded_file) -> str:
    """Save an uploaded file and return its path."""
    try:
        file_path = os.path.join('uploads', uploaded_file.name)
        with open(file_path, 'wb') as f:
            f.write(uploaded_file.getvalue())
        return file_path
    except Exception as e:
        st.error(f"Error saving file {uploaded_file.name}: {str(e)}")
        return None

def load_index_log():
    if not os.path.exists(INDEX_LOG):
        return {}
    try:
        with open(INDEX_LOG, "r") as f:
            return json.load(f)
    except Exception:
        return {}

def save_index_log(log):
    with open(INDEX_LOG, "w") as f:
        json.dump(log, f, indent=2)

def mark_file_indexed(filename, remove=False):
    """Mark a file as indexed or remove it from the index log."""
    log = load_index_log()
    if remove:
        log.pop(filename, None)
    else:
        log[filename] = {"timestamp": datetime.now().isoformat()}
    save_index_log(log)

def get_unprocessed_files():
    log = load_index_log()
    files = os.listdir(UPLOAD_DIR)
    return [f for f in files if f not in log]

def get_processed_files():
    log = load_index_log()
    return list(log.keys())

def move_to_processed(file_path):
    dest_path = os.path.join(PROCESSED_DIR, os.path.basename(file_path))
    shutil.move(file_path, dest_path)

def get_chroma_client():
    """Get or create ChromaDB client with proper settings."""
    try:
        client = chromadb.HttpClient(host=os.getenv("CHROMA_HOST", "localhost"), port=int(os.getenv("CHROMA_PORT", 8000)))
        
        # Test persistent mode
        print("[LOG] Testing persistent mode...")
        # Clean up test collection if it already exists
        try:
            client.delete_collection("test_collection")
        except Exception:
            pass  # Ignore if it doesn't exist

        test_collection = client.create_collection("test_collection")
        test_collection.add(
            documents=["test"],
            metadatas=[{"source": "test"}],
            ids=["test"]
        )
        test_collection.get(ids=["test"])
        client.delete_collection("test_collection")
        print("[LOG] Persistent mode test successful")
        
        return client
    except Exception as e:
        st.error(f"Error initializing ChromaDB: {str(e)}")
        print(f"Error initializing ChromaDB: {str(e)}")
        raise

@st.cache_resource(show_spinner="Loading embedding model...")
def load_embedding_model():
    MODEL_NAME = 'BAAI/bge-base-en-v1.5'  # or 'sentence-transformers/all-mpnet-base-v2'
    model = SentenceTransformer(MODEL_NAME)
    is_bge = MODEL_NAME.startswith('BAAI/bge') or MODEL_NAME.startswith('intfloat/e5')
    return model, is_bge

st_model, IS_BGE = load_embedding_model()

class StreamlitEmbeddingFunction:
    def __init__(self, model, is_bge):
        self.model = model
        self.is_bge = is_bge
    def __call__(self, input):
        if self.is_bge:
            input = [f"passage: {t}" for t in input]
        result = self.model.encode(input, show_progress_bar=False)
        # Handle both numpy arrays and lists
        if hasattr(result, 'tolist'):
            return result.tolist()
        elif isinstance(result, list):
            return result
        else:
            return result.tolist() if hasattr(result, 'tolist') else list(result)
    @staticmethod
    def name():
        return "streamlit-sentence-transformers"

st_embedding_fn = StreamlitEmbeddingFunction(st_model, IS_BGE)

def get_chroma_collection():
    """Get or create ChromaDB collection with proper settings."""
    try:
        client = get_chroma_client()
        # Only get or create, do not delete!
        collection = client.get_or_create_collection(
            name="legal_documents",
            embedding_function=st_embedding_fn
        )
        return collection
    except Exception as e:
        st.error(f"Error getting ChromaDB collection: {str(e)}")
        print(f"Error getting ChromaDB collection: {str(e)}")
        raise

def add_log(msg):
    st.session_state.log_buffer.append(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")
    if len(st.session_state.log_buffer) > LOG_BUFFER_SIZE:
        st.session_state.log_buffer = st.session_state.log_buffer[-LOG_BUFFER_SIZE:]
    print(msg)

def update_doc_status(filename, status, chunk_count=None):
    st.session_state.doc_status[filename] = {
        'status': status,
        'chunk_count': chunk_count,
        'last_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }

def get_total_chunks():
    try:
        client = get_chroma_client()
        collection = get_chroma_collection()
        return collection.count()
    except:
        return 0

def get_storage_usage():
    total = 0
    for dirpath, dirnames, filenames in os.walk('storage'):
        for f in filenames:
            fp = os.path.join(dirpath, f)
            total += os.path.getsize(fp)
    return total // 1024  # KB

def load_documents():
    """Load and process only unprocessed documents from the uploads directory."""
    try:
        st.session_state.processing = True
        add_log("Starting document processing...")
        
        # Get unprocessed files
        file_list = get_unprocessed_files()
        add_log(f"Found {len(file_list)} unprocessed files")
        
        if not file_list:
            st.info("No new documents to process. All files in uploads/ have been indexed.")
            st.session_state.processing = False
            return None
            
        # Initialize ChromaDB
        try:
            client = get_chroma_client()
            collection = get_chroma_collection()
            add_log("Successfully initialized ChromaDB")
        except Exception as e:
            st.error(f"Failed to initialize ChromaDB: {str(e)}")
            add_log(f"Failed to initialize ChromaDB: {str(e)}")
            st.session_state.processing = False
            return None
            
        # Create a status container at the top
        status_container = st.empty()
        progress_bar = st.progress(0, text="Processing documents...")
        total_files = len(file_list)
        
        # Live log container that only shows during processing
        log_container = st.container()
        
        for file_idx, filename in enumerate(file_list):
            file_path = os.path.join(UPLOAD_DIR, filename)
            add_log(f"Processing file {file_idx+1}/{total_files}: {filename}")
            update_doc_status(filename, 'Processing')
            
            # Update status container
            status_container.markdown(f"""
            ### Processing Status
            - **Current File:** {filename}
            - **Progress:** {file_idx + 1}/{total_files} files
            - **Status:** Processing
            """)
            
            # Update live logs during processing
            with log_container:
                st.text("\n".join(st.session_state.log_buffer[-5:]))
            
            log_memory_usage(f"Before processing {filename}")
            chunks = process_document(file_path)
            
            if not chunks:
                st.warning(f"No chunks extracted from {filename}")
                update_doc_status(filename, 'Error', 0)
                add_log(f"No chunks extracted from {filename}")
                continue
                
            for i, chunk in enumerate(chunks):
                try:
                    add_log(f"Processing chunk {i+1}/{len(chunks)} of {filename}")
                    # Update live logs during processing
                    with log_container:
                        st.text("\n".join(st.session_state.log_buffer[-5:]))
                        
                    embedding = st_embedding_fn([chunk['text']])[0]
                    collection.add(
                        embeddings=[embedding],
                        documents=[chunk['text']],
                        metadatas=[chunk['metadata']],
                        ids=[f"{filename}_{chunk['metadata']['chunk_id']}"]
                    )
                    
                except Exception as e:
                    st.error(f"Error processing chunk {i+1} of {filename}: {str(e)}")
                    add_log(f"Error processing chunk {i+1} of {filename}: {str(e)}")
                    continue
                    
                progress_bar.progress((file_idx + (i+1)/len(chunks)) / total_files, 
                                    text=f"Processing {filename} ({i+1}/{len(chunks)})")
                                    
            update_doc_status(filename, 'Indexed', len(chunks))
            mark_file_indexed(filename)
            move_to_processed(file_path)
            gc.collect()
            progress_bar.progress((file_idx+1)/total_files, text=f"Processed {filename}")
            
        # Update final status
        status_container.markdown("""
        ### Processing Complete! ✅
        - All documents have been processed and indexed
        - You can now search your documents
        """)
        
        # Clear the live log container after processing
        log_container.empty()
        
        st.session_state.documents_loaded = True
        st.success("Document processing completed!")
        add_log("Document processing completed!")
        st.session_state.processing = False
        return collection
        
    except Exception as e:
        st.error(f"Error loading documents: {str(e)}")
        add_log(f"Error loading documents: {str(e)}")
        st.session_state.processing_error = str(e)
        st.session_state.processing = False
        return None

def search_documents(query: str, collection) -> List[Dict[str, Any]]:
    """Search documents using ChromaDB."""
    try:
        if not query.strip():
            return []
            
        # For BGE/E5, prefix 'query: ' to the query
        if IS_BGE:
            query_for_embedding = f"query: {query}"
        else:
            query_for_embedding = query
            
        # Get the embedding for the query
        query_embedding = st_embedding_fn([query_for_embedding])[0]
        
        # Use the embedding directly in the query
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=5,
            include=["documents", "metadatas", "distances"]
        )
        
        # Format results with better metadata
        formatted_results = []
        for i, (doc, metadata, distance) in enumerate(zip(
            results["documents"][0],
            results["metadatas"][0],
            results["distances"][0]
        )):
            formatted_results.append({
                "text": doc,
                "source": metadata.get("filename", "Unknown"),
                "chunk": metadata.get("chunk_id", 0),
                "distance": float(distance),  # Ensure distance is a float
                "rank": i + 1
            })
        return formatted_results
    except Exception as e:
        st.error(f"Error searching documents: {str(e)}")
        print(f"Search error details: {str(e)}")  # Add detailed error logging
        return []

# Create necessary directories
os.makedirs('storage', exist_ok=True)
os.makedirs('uploads', exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)

# Streamlit UI
st.title("Legal Document RAG System")

# Sidebar for document upload
with st.sidebar:
    st.header("📁 Document Upload")
    
    # Collapsible System Monitor
    with st.expander("System Monitor", expanded=False):
        cpu = psutil.cpu_percent(interval=0.1)
        mem = psutil.virtual_memory().percent
        st.markdown(f"**CPU Usage:** {cpu}%")
        st.markdown(f"**Memory Usage:** {mem}%")
        st.markdown(f"**Documents Indexed:** {len(get_processed_files())}")
        st.markdown(f"**Total Chunks:** {get_total_chunks()}")
        st.markdown(f"**Storage Usage:** {get_storage_usage()} KB")

    # Upload UI - simplified
    uploaded_files = st.file_uploader(
        "Upload Legal Documents",
        type=['pdf', 'txt', 'docx', 'md'],
        accept_multiple_files=True,
        help="Upload your case documents. Supported formats: PDF, TXT, DOCX, MD"
    )
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            if uploaded_file.name not in st.session_state.uploaded_files:
                with st.spinner(f"Uploading {uploaded_file.name}..."):
                    try:
                        file_path = save_uploaded_file(uploaded_file)
                        if file_path:
                            st.session_state.uploaded_files.append(uploaded_file.name)
                            st.success(f"✅ {uploaded_file.name}")
                            st.rerun()
                    except Exception as e:
                        st.error(f"❌ Error uploading {uploaded_file.name}: {str(e)}")

    # Compact file list
    if st.session_state.uploaded_files:
        st.subheader("📄 Uploaded Files")
        for file_name in st.session_state.uploaded_files[:5]:  # Show only first 5
            if os.path.exists(os.path.join('uploads', file_name)):
                st.write(f"• {file_name}")
            else:
                st.session_state.uploaded_files.remove(file_name)
                st.rerun()
        
        if len(st.session_state.uploaded_files) > 5:
            st.write(f"... and {len(st.session_state.uploaded_files) - 5} more files")

    # Collapsible Document Management
    with st.expander("🔧 Document Management", expanded=False):
        # Process Documents Button
        unprocessed_count = len(get_unprocessed_files())
        if unprocessed_count > 0:
            if st.button(f"Process {unprocessed_count} Documents", type="primary"):
                st.session_state.show_logs = True
                st.rerun()
        
        # Document Status
        if st.session_state.doc_status:
            st.write("**Document Status:**")
            for filename, status in st.session_state.doc_status.items():
                status_emoji = {"Processing": "⏳", "Indexed": "✅", "Error": "❌", "Queued": "📋"}.get(status['status'], "❓")
                st.write(f"{status_emoji} {filename[:20]}{'...' if len(filename) > 20 else ''}")
        
        # Reset button moved here
        if st.button("🗑️ Reset All Data", type="secondary"):
            try:
                # Delete uploads, processed, storage, and index log
                for folder in ['uploads', PROCESSED_DIR, 'storage']:
                    if os.path.exists(folder):
                        shutil.rmtree(folder)
                    os.makedirs(folder, exist_ok=True)
                    os.chmod(folder, 0o777)
                # Remove index log
                if os.path.exists(INDEX_LOG):
                    os.remove(INDEX_LOG)
                # Also delete ChromaDB collection
                try:
                    client = chromadb.HttpClient(host=os.getenv("CHROMA_HOST", "localhost"), port=int(os.getenv("CHROMA_PORT", 8000)))
                    client.delete_collection("legal_documents")
                except Exception as e:
                    st.warning(f"Could not delete ChromaDB collection: {e}")
                # Clear session state
                st.session_state.last_processed_hash = None
                st.session_state.documents_loaded = False
                st.session_state.uploaded_files = []
                st.session_state.doc_status = {}
                st.session_state.log_buffer = []
                st.session_state.show_logs = False
                st.session_state.llm_responses = []
                # Force garbage collection
                gc.collect()
                st.success("All data reset!")
                st.rerun()
            except Exception as e:
                st.error(f"Error resetting: {str(e)}")

# Main content area
# Process documents if needed
if check_documents_changed():
    with st.spinner("Updating document index..."):
        try:
            # Reset ChromaDB
            if os.path.exists('storage'):
                shutil.rmtree('storage')
            os.makedirs('storage', exist_ok=True)
            os.chmod('storage', 0o777)
            
            # Process documents
            st.session_state.show_logs = True
            collection = load_documents()
            if collection:
                st.success("Document index updated successfully!")
                st.session_state.show_logs = False
                gc.collect()
            else:
                st.error("Failed to process documents.")
        except Exception as e:
            st.error(f"Error updating document index: {str(e)}")
            st.session_state.processing_error = str(e)

# Show logs only when explicitly requested or during processing
if st.session_state.show_logs and st.session_state.log_buffer:
    with st.expander("📋 Processing Logs", expanded=st.session_state.processing):
        if st.button("Hide Logs"):
            st.session_state.show_logs = False
            st.rerun()
        
        # Show recent logs
        recent_logs = st.session_state.log_buffer[-10:]
        st.text("\n".join(recent_logs))
        
        if st.button("Clear Logs"):
            st.session_state.log_buffer = []
            st.session_state.show_logs = False
            st.rerun()

# Query section
st.header("🔍 Document Query")

# Compact status display
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("Documents", len(get_processed_files()))
with col2:
    st.metric("Chunks", get_total_chunks())
with col3:
    storage_mb = get_storage_usage() / 1024
    st.metric("Storage", f"{storage_mb:.1f} MB")

# Query input
query = st.text_area(
    "Enter your query",
    height=100,
    placeholder="Ask a question about your legal documents...",
    help="Ask a question about your legal documents"
)

# Search functionality
processed_files = get_processed_files()

if st.button("🔍 Search", type="primary"):
    if not processed_files:
        st.error("Please upload and process some documents first!")
    elif not query.strip():
        st.warning("Please enter a query.")
    else:
        with st.spinner("Searching documents..."):
            try:
                client = get_chroma_client()
                collection = get_chroma_collection()
                results = search_documents(query, collection)
                
                if results:
                    st.subheader("Search Results")
                    # Model-agnostic entity aggregation and display (using model attributes)
                    aggregated_entities = defaultdict(list)
                    st.session_state.llm_responses = []  # Clear previous responses
                    for result in results:
                        entities = extract_entities(result['text'])
                        for entity_type in entities.__fields__:
                            parties = getattr(entities, entity_type)
                            aggregated_entities[entity_type].extend([
                                p for p in parties if getattr(p, 'source', None) == 'llm'
                            ])
                    
                    # Display LLM Raw Responses
                    st.subheader("🤖 LLM Raw Responses")
                    if st.session_state.llm_responses:
                        for response in st.session_state.llm_responses:
                            with st.expander(f"Response for {response['entity_type']}", expanded=True):
                                st.code(response['content'], language='json')
                    else:
                        st.info("No LLM responses available.")
                    
                    st.markdown("---")  # Add a separator
                    
                    st.subheader("📋 Extracted Legal Entities")
                    if not any(aggregated_entities.values()):
                        st.info("No legal entities found in the search results.")
                    else:
                        for entity_type, parties in aggregated_entities.items():
                            if parties:
                                label = entity_type.replace('_', ' ').capitalize()
                                st.markdown(f"**{label}:**")
                                for party in parties:
                                    role_display = f" ({party.role})" if party.role and party.role.lower() not in entity_type else ""
                                    st.write(f"• {party.name}{role_display}")
                    st.markdown("---")  # Add a separator
                    
                    # Show the raw results as before
                    for i, result in enumerate(results, 1):
                        relevance = 1 - result['distance']
                        with st.expander(f"Result {i} - {result['source']} (Relevance: {relevance:.1%})", 
                                       expanded=i==1):  # Expand first result by default
                            st.write(f"**Source:** {result['source']}")
                            st.write(f"**Chunk:** {result['chunk']}")
                            st.write(f"**Relevance Score:** {relevance:.1%}")
                            st.write("**Content:**")
                            st.write(result['text'])
                else:
                    st.info("No relevant results found. Try rephrasing your query.")
            except Exception as e:
                st.error(f"Error during search: {str(e)}")
            finally:
                gc.collect()

# Help section - collapsed by default
with st.expander("ℹ️ How to Use", expanded=False):
    st.markdown("""
    ### Quick Start
    1. **Upload documents** using the sidebar file uploader
    2. **Process documents** using the "Process Documents" button in Document Management
    3. **Enter your query** in the text area above
    4. **Click Search** to find relevant information

    ### Supported Document Types
    - PDF files (.pdf)
    - Text files (.txt)
    - Word documents (.docx)
    - Markdown files (.md)

    ### Tips
    - Keep document sizes under 10MB for best performance
    - Use specific keywords in your queries for better results
    - Check the Document Management section to monitor processing status
    """)